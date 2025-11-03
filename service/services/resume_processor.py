import PyPDF2
import docx
import re
from typing import Dict

def extract_text_from_pdf(file_path: str) -> str:
    """Extract all text from PDF"""
    text = ""
    try:
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
    except Exception as e:
        print(f"Error reading PDF: {e}")
    return text

def extract_text_from_docx(file_path: str) -> str:
    """Extract all text from DOCX"""
    text = ""
    try:
        doc = docx.Document(file_path)
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
            text += "\n"
    except Exception as e:
        print(f"Error reading DOCX: {e}")
    return text

def extract_email(text: str) -> str:
    """Extract email using regex"""
    email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
    emails = re.findall(email_pattern, text)
    return emails[0] if emails else ""

def extract_phone(text: str) -> str:
    """Extract phone number"""
    # Indian phone patterns
    phone_patterns = [
        r'\+91[\s-]?\d{10}',  # +91 followed by 10 digits
        r'\d{10}',  # 10 digits
        r'\d{3}[-.\s]?\d{3}[-.\s]?\d{4}',  # US format
    ]
    
    for pattern in phone_patterns:
        phones = re.findall(pattern, text)
        if phones:
            return phones[0]
    return ""

def extract_skills(text: str) -> str:
    """Extract technical skills with better accuracy"""
    skills_keywords = [
        # Programming Languages (add more variations)
        'python', 'java', 'javascript', 'js', 'typescript', 'ts', 'c++', 'cpp', 'c#', 'csharp',
        'golang', 'go', 'rust', 'kotlin', 'swift', 'php', 'ruby', 'scala', 'r programming', 'matlab',
        
        # Frameworks & Libraries
        'react', 'reactjs', 'angular', 'vue', 'vuejs', 'nodejs', 'node.js', 'node', 'express', 
        'expressjs', 'django', 'flask', 'fastapi', 'spring boot', 'spring', 'laravel', 
        'nextjs', 'next.js', 'nuxtjs', 'svelte',
        
        # Frontend
        'html', 'html5', 'css', 'css3', 'sass', 'scss', 'less', 'tailwind', 'tailwindcss',
        'bootstrap', 'jquery', 'webpack', 'vite', 'redux', 'mobx',
        
        # Mobile
        'react native', 'flutter', 'android', 'ios', 'swift', 'kotlin', 'xamarin', 'ionic',
        
        # Databases
        'sql', 'mysql', 'postgresql', 'postgres', 'mongodb', 'mongo', 'redis', 'elasticsearch',
        'dynamodb', 'cassandra', 'oracle', 'sqlite', 'mariadb', 'neo4j', 'couchdb',
        
        # Cloud & DevOps
        'aws', 'amazon web services', 'azure', 'microsoft azure', 'gcp', 'google cloud', 
        'docker', 'kubernetes', 'k8s', 'jenkins', 'gitlab', 'github actions', 'circleci',
        'terraform', 'ansible', 'puppet', 'chef', 'ci/cd', 'devops', 'vagrant',
        
        # AI/ML/Data Science
        'machine learning', 'ml', 'deep learning', 'dl', 'tensorflow', 'pytorch', 'keras',
        'scikit-learn', 'sklearn', 'pandas', 'numpy', 'scipy', 'nlp', 'natural language processing',
        'computer vision', 'opencv', 'ai', 'artificial intelligence', 'data science',
        'neural networks', 'cnn', 'rnn', 'lstm', 'transformers', 'bert', 'gpt',
        
        # Testing
        'testing', 'unit testing', 'integration testing', 'junit', 'pytest', 'jest',
        'mocha', 'chai', 'selenium', 'cypress', 'playwright', 'testng',
        
        # Version Control & Tools
        'git', 'github', 'gitlab', 'bitbucket', 'svn', 'mercurial',
        
        # APIs & Protocols
        'rest api', 'restful', 'graphql', 'soap', 'grpc', 'websocket', 'mqtt',
        
        # Architecture & Patterns
        'microservices', 'monolithic', 'mvc', 'mvvm', 'serverless', 'event-driven',
        
        # Methodologies
        'agile', 'scrum', 'kanban', 'waterfall', 'tdd', 'bdd', 'lean',
        
        # Other Technologies
        'kafka', 'rabbitmq', 'apache', 'nginx', 'linux', 'ubuntu', 'centos',
        'bash', 'shell scripting', 'powershell', 'jira', 'confluence',
        'blockchain', 'ethereum', 'solidity', 'web3', 'solana',
        'sap', 'erp', 'crm', 'salesforce', 'seo', 'sem'
    ]
    
    text_lower = text.lower()
    found_skills = set()
    
    # Use more sophisticated matching
    for skill in skills_keywords:
        # Match whole words or common variations
        patterns = [
            r'\b' + re.escape(skill) + r'\b',
            r'\b' + re.escape(skill.replace(' ', '')) + r'\b',  # Remove spaces
            r'\b' + re.escape(skill.replace('-', '')) + r'\b',  # Remove hyphens
        ]
        
        for pattern in patterns:
            if re.search(pattern, text_lower):
                # Use original casing from text if possible
                found_skills.add(skill.title())
                break
    
    return ", ".join(sorted(found_skills)) if found_skills else "Not specified"


def extract_experience(text: str) -> str:
    """Extract work experience with better section detection"""
    text_lower = text.lower()
    
    # Multiple keywords for experience section
    experience_keywords = [
        'professional experience', 'work experience', 'experience',
        'employment history', 'work history', 'career history',
        'professional background'
    ]
    
    # Try to find the experience section
    for keyword in experience_keywords:
        pattern = r'\b' + re.escape(keyword) + r'\b'
        match = re.search(pattern, text_lower)
        
        if match:
            start_idx = match.start()
            
            # Find where experience section ends (look for next major section)
            end_keywords = ['education', 'skills', 'certifications', 'projects', 'achievements']
            end_idx = len(text)
            
            for end_keyword in end_keywords:
                end_pattern = r'\b' + re.escape(end_keyword) + r'\b'
                end_match = re.search(end_pattern, text_lower[start_idx + 100:])  # Skip first 100 chars
                if end_match:
                    potential_end = start_idx + 100 + end_match.start()
                    end_idx = min(end_idx, potential_end)
            
            experience_text = text[start_idx:end_idx]
            
            # Limit to reasonable length
            if len(experience_text) > 2000:
                experience_text = experience_text[:2000] + "..."
            
            return experience_text.strip()
    
    # Fallback: look for company names and dates
    # Pattern: Company Name (Year-Year) or (Month Year - Month Year)
    date_patterns = [
        r'[A-Z][a-z\s&,\.]+(?:\([0-9]{4}\s*-\s*(?:[0-9]{4}|Present|Current)\))',
        r'[A-Z][a-z\s&,\.]+\n.*?[0-9]{4}\s*-\s*(?:[0-9]{4}|Present|Current)'
    ]
    
    experience_matches = []
    for pattern in date_patterns:
        matches = re.findall(pattern, text, re.MULTILINE)
        experience_matches.extend(matches)
    
    if experience_matches:
        return " | ".join(experience_matches[:5])  # Top 5 experiences
    
    return "Not specified"


def extract_education(text: str) -> str:
    """Extract education with better accuracy"""
    education_patterns = [
        # Degree patterns
        r'(B\.?Tech|Bachelor of Technology|B\.?E\.?|Bachelor of Engineering)\s+in\s+[\w\s]+',
        r'(M\.?Tech|Master of Technology|M\.?E\.?|Master of Engineering)\s+in\s+[\w\s]+',
        r'(MBA|Master of Business Administration)',
        r'(Ph\.?D\.?|Doctor of Philosophy)\s+in\s+[\w\s]+',
        r'(B\.?Sc\.?|Bachelor of Science)\s+in\s+[\w\s]+',
        r'(M\.?Sc\.?|Master of Science)\s+in\s+[\w\s]+',
        r'(BCA|Bachelor of Computer Applications)',
        r'(MCA|Master of Computer Applications)',
        
        # Institution patterns
        r'(IIT|Indian Institute of Technology)\s+[\w\s]+',
        r'(NIT|National Institute of Technology)\s+[\w\s]+',
        r'(IIIT|Indian Institute of Information Technology)\s+[\w\s]+',
        r'University\s+of\s+[\w\s]+',
        r'[\w\s]+\s+University',
        r'[\w\s]+\s+College'
    ]
    
    text_clean = text.replace('\n', ' ')
    education_parts = []
    
    for pattern in education_patterns:
        matches = re.findall(pattern, text_clean, re.IGNORECASE)
        for match in matches:
            if isinstance(match, tuple):
                match = ' '.join(match)
            # Get context around match
            match_pos = text_clean.lower().find(match.lower())
            if match_pos != -1:
                context = text_clean[max(0, match_pos-50):min(len(text_clean), match_pos+len(match)+100)]
                education_parts.append(context.strip())
    
    if education_parts:
        # Remove duplicates while preserving order
        seen = set()
        unique_edu = []
        for edu in education_parts:
            edu_lower = edu.lower()
            if edu_lower not in seen:
                seen.add(edu_lower)
                unique_edu.append(edu)
        
        return " | ".join(unique_edu[:3])  # Top 3 education entries
    
    return "Not specified"


def extract_years_of_experience(text: str) -> str:
    """Extract years of experience more accurately"""
    patterns = [
        r'(\d+)\+?\s*years?\s+of\s+(?:professional\s+)?experience',
        r'(?:professional\s+)?experience\s*:?\s*(\d+)\+?\s*years?',
        r'(\d+)\+?\s*yrs?\s+(?:of\s+)?experience',
        r'total\s+experience\s*:?\s*(\d+)\+?\s*years?',
    ]
    
    text_lower = text.lower()
    
    for pattern in patterns:
        matches = re.findall(pattern, text_lower)
        if matches:
            years = matches[0]
            return f"{years} years"
    
    # Try to calculate from experience section
    # Look for date ranges like "2018 - 2023" or "Jan 2018 - Present"
    year_ranges = re.findall(r'(\d{4})\s*-\s*(?:(\d{4})|present|current)', text_lower)
    
    if year_ranges:
        total_years = 0
        current_year = 2025  # Update this year as needed
        
        for start, end in year_ranges:
            start_year = int(start)
            end_year = int(end) if end else current_year
            years = end_year - start_year
            total_years += max(0, years)
        
        if total_years > 0:
            return f"{total_years} years"
    
    return "Not specified"

def process_resume(file_path: str) -> Dict:
    """
    Complete OCR processing of resume
    Extract all possible information
    """
    # Extract text based on file type
    if file_path.endswith('.pdf'):
        raw_text = extract_text_from_pdf(file_path)
    elif file_path.endswith('.docx'):
        raw_text = extract_text_from_docx(file_path)
    else:
        raise ValueError("Unsupported file format. Only PDF and DOCX allowed.")
    
    if not raw_text or len(raw_text.strip()) < 50:
        raise ValueError("Could not extract meaningful text from resume")
    
    # Extract all information
    extracted_data = {
        "raw_text": raw_text.strip(),
        "email": extract_email(raw_text),
        "phone": extract_phone(raw_text),
        "skills": extract_skills(raw_text),
        "experience": extract_experience(raw_text),
        "education": extract_education(raw_text),
        "years_of_experience": extract_years_of_experience(raw_text),
    }
    
    return extracted_data